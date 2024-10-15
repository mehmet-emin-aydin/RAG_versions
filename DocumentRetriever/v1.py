import streamlit as st
import PyPDF2
from langchain.chains import StuffDocumentsChain,MapReduceDocumentsChain, ReduceDocumentsChain
from langchain_text_splitters import CharacterTextSplitter, RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from langchain_openai import AzureChatOpenAI
from langchain.chains.llm import LLMChain
from langchain_core.prompts import PromptTemplate
from langchain.docstore.document import Document
import pandas as pd
from datetime import datetime
import io
import docx
import os
from dotenv import load_dotenv

load_dotenv()
log_data = [] 
import sys
sys.path.append('../..')

azure_oai_key = os.getenv("AZURE_OPENAI_API_KEY")
azure_oai_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
api_version = "2024-02-15-preview"
llm_name = "gpt-4o"


def chunk_text(text: str) -> list[str]:
    chunks = None
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=4000,
        chunk_overlap=0,
        separators=["\n\n", "\n", " ", ""]
    )
    chunks = text_splitter.split_text(text)
    return chunks

def process_uploaded_files(uploaded_files):

# if uploaded_file is not None:
#     docs = []
#     reader = PdfReader(uploaded_file)
#     i = 1
#     for page in reader.pages:
#         docs.append(Document(page_content=page.extract_text(), metadata={'page':i}))
#         i += 1

    docs = []
    for file in uploaded_files:
        file_name = file.name
        file_extension = os.path.splitext(file_name)[1]
        if file_extension == '.txt':
            text = file.read().decode('utf-8')
            chunks = chunk_text(text)
            [docs.extend(Document(page_content=chunks[i], metadata = {'source':file.name, 'page': i}) for i in range(len(chunks)))]
        elif file_extension == '.pdf':
            reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
            [docs.extend(Document(page_content=page.extract_text(), metadata = {'source':file.name, 'page': i}) for i, page in enumerate(reader.pages))]
        elif file_extension == '.docx':
            doc = docx.Document(io.BytesIO(file.read()))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            chunks = chunk_text(text)
            [docs.extend(Document(page_content=chunks[i], metadata = {'source':file.name, 'page': i}) for i in range(len(chunks)))]

 

    # First, we specify the LLMChain to use for mapping each document to an individual summary:
    llm=AzureChatOpenAI(model="gpt-4o", temperature=0, api_key=azure_oai_key, azure_endpoint=azure_oai_endpoint, api_version=api_version)

    # Map
    map_template = """The following is a set of documents
    {docs}
    Based on this list of docs, please identify the main themes 
    Helpful Answer in the language of the provided documents:"""
    map_prompt = PromptTemplate.from_template(map_template)
    map_chain = LLMChain(llm=llm, prompt=map_prompt)


    # Reduce
    reduce_template = """The following is set of summaries:
    {docs}
    Take these and distill it into a final, consolidated summary of the main themes. 
    Helpful Answer in the language of the provided documents::"""
    reduce_prompt = PromptTemplate.from_template(reduce_template)


    # Run chain
    reduce_chain = LLMChain(llm=llm, prompt=reduce_prompt)

    # Takes a list of documents, combines them into a single string, and passes this to an LLMChain
    combine_documents_chain = StuffDocumentsChain(
        llm_chain=reduce_chain, document_variable_name="docs"
    )

    # Combines and iteratively reduces the mapped documents
    reduce_documents_chain = ReduceDocumentsChain(
        # This is final chain that is called.
        combine_documents_chain=combine_documents_chain,
        # If documents exceed context for `StuffDocumentsChain`
        collapse_documents_chain=combine_documents_chain,
        # The maximum number of tokens to group documents into.
        token_max=10000,
    )


    # Combining documents by mapping a chain over them, then combining results
    map_reduce_chain = MapReduceDocumentsChain(
        # Map chain
        llm_chain=map_chain,
        # Reduce chain
        reduce_documents_chain=reduce_documents_chain,
        # The variable name in the llm_chain to put the documents in
        document_variable_name="docs",
        # Return the results of the map steps in the output
        return_intermediate_steps=False,
    )

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=4000,
        chunk_overlap=100,
        separators=["\n\n", "\n", " ", ""]
    )
    split_docs = text_splitter.split_documents(documents=docs)
    print(len(split_docs))
    result = map_reduce_chain.invoke(split_docs)

    return result["output_text"]
def main():

    st.set_page_config(page_title="Document Summarizer")
    st.title("Document Summarizing App")
    st.write("Summarize your pdf files in just a few seconds.")
    # Displaying a description
    st.divider() # Inserting a divider for better layout
    # Creating a file uploader widget to upload PDF files
    pdf = st. file_uploader ('Upload your PDF Document', accept_multiple_files=True)
    # Creating a button for users to submit their PDF
    if pdf is not None:
        submit = st.button("Generate Summary")
        if submit:
            summary = process_uploaded_files(pdf)

            txt = f'Source Document: {" ".join(file.name for file in pdf)} \n Summary:\n {summary}'
            # df = pd.DataFrame(data)
            # txt = df.to_csv().encode('utf-8')

            datetime_stamp = f"{datetime.now().strftime('%Y-%m-%d-%H-%M-%S')}"

            st.download_button(
                label="Download Summary as Text File",
                data=txt,
                file_name=f'Summarization_{datetime_stamp}.txt'
            )
            st.subheader("Summary of the file:")
            st.write(summary)
if __name__ == '__main__':
    main()