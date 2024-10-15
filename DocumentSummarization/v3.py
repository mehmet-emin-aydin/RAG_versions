import streamlit as st
import PyPDF2
from langchain.chains import StuffDocumentsChain,MapReduceDocumentsChain, ReduceDocumentsChain
from langchain_text_splitters import CharacterTextSplitter, RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from langchain_openai import AzureChatOpenAI
from langchain.chains.llm import LLMChain
from langchain_core.prompts import PromptTemplate
from langchain.docstore.document import Document
import docx
import os
import io
from dotenv import load_dotenv

load_dotenv()
log_data = [] 
import sys
sys.path.append('../..')

azure_oai_key = os.getenv("AZURE_OPENAI_API_KEY")
azure_oai_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
api_version = "2024-02-15-preview"
llm_name = "gpt-4o"

def chunk_text(text: str) -> list[str]:
    chunks = None
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=100,
        separators=["\n\n", "\n", " ", ""]
    )
    chunks = text_splitter.split_text(text)
    print(f'first chunk:{chunks[0]}')
    return chunks

def process_uploaded_files(uploaded_files):
    # pdf_path = '/home/mehmet/Desktop/Softtech/RAG/RAG_versions/test_files/rag.pdf'
    # loader = PyPDFLoader(file_path=pdf_path)
    # docs = loader.load()
    docs = []
    for file in uploaded_files:
        doc = docx.Document(io.BytesIO(file.read()))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        chunks = chunk_text(text)
        [docs.extend(Document(page_content=chunks[i], metadata = {'source':file.name, 'page': i}) for i in range(len(chunks)))]
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=100,
        separators=["\n\n", "\n"," ", ""]
    )
    split_docs = text_splitter.split_documents(docs)

    return {'docs': doc.paragraphs[0].text, 'split':split_docs}





def main():

    st.set_page_config(page_title="PDF Summarizer")
    st.title("PDF Summarazing App")
    st.write("Summarize your pdf files in just a few seconds.")
    # Displaying a description
    st.divider() # Inserting a divider for better layout
    # Creating a file uploader widget to upload PDF files
    pdf = st. file_uploader ('Upload your PDF Document', accept_multiple_files=True)
    # Creating a button for users to submit their PDF

    submit = st.button("Generate Summary")
    if submit and pdf is not None:
        summary = process_uploaded_files(pdf)
        st.subheader("Summary of the file:")
        st.write(summary)


if __name__ == '__main__':
    main()